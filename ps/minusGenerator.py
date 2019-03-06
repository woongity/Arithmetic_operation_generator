from docx import Document
import random

document = Document()

document.add_heading("곱하기 연습문제", 0)
document.add_heading("두자리수끼리 빼기 연습문제", 1)
count = 0
while count < 50:
    a = random.randint(10, 99)
    b = random.randint(10, 99)
    if a > b:
        count += 1
    else:
        continue
    problem = str(a) + "-" + str(b) + " = "
    document.add_paragraph(problem)
document.add_heading("세자리수끼리 곱하기 연습문제", 1)
while count < 100:
    a = random.randint(100, 999)
    b = random.randint(100, 999)
    if a > b:
        count += 1
    else:
        continue
    problem = str(a) + "-" + str(b) + " = "
    document.add_paragraph(problem)
document.add_heading("네자리수끼리 곱하기 연습문제", 1)
while count < 150:
    a = random.randint(1000, 9999)
    b = random.randint(1000, 9999)
    if a > b:
        count += 1
    else:
        continue
    problem = str(a) + "-" + str(b) + " = "
    document.add_paragraph(problem)

document.save('minus.docx')
