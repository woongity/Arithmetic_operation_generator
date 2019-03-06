from docx import Document
import random
document=Document()

document.add_heading("곱하기 연습문제",0)
document.add_heading("두자리수끼리 더하기 연습문제",1)
for multiples in range(1,50):
    a=random.randint(10,99)
    b=random.randint(10,99)
    problem=str(a)+"+"+str(b)+" = "
    document.add_paragraph(problem)
document.add_heading("세자리수끼리 곱하기 연습문제",1)
for multiples in range(1,50):
    a=random.randint(100,999)
    b=random.randint(100,999)
    problem=str(a)+"+"+str(b)+" = "
    document.add_paragraph(problem)
for multiples in range(1,50):
    a=random.randint(1000,9999)
    b=random.randint(1000,9999)
    problem=str(a)+"+"+str(b)+" = "
    document.add_paragraph(problem)

document.save('multiple.docx')