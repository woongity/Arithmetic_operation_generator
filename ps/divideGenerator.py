from docx import Document
import random
document=Document()

document.add_heading("나누기 연습문제",0)

document.add_heading("소수점 둘째 자리까지 구하시오",3)

document.add_heading("두자리수 나누기 한자리수",1)
count = 0
while count < 50:
    a=random.randint(2,9)
    b=random.randint(1,9)
    if a > b:
        count += 1
    else:
        continue

    problem=str(a)+"÷"+str(b)+" = "
    document.add_paragraph(problem)
document.add_heading("두자리수 나누기 두자리수" , 1)
while count<100:
    a=random.randint(10,99)
    b=random.randint(10,99)
    if a > b:
        count += 1
    else:
        continue
    problem=str(a)+"÷"+str(b)+" = "
    document.add_paragraph(problem)
document.add_heading("세자리수끼리 나누기 연습문제" , 1)
while count<100:
    a=random.randint(100,999)
    b=random.randint(100,999)
    if a > b:
        count += 1
    else:
        continue

    problem=str(a)+"÷"+str(b)+" = "
    document.add_paragraph(problem)

document.save('divide.docx')